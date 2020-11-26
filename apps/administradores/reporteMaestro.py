from django.http import HttpResponse
# Manejo Word
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Inches
# Modelos
from apps.reportes.models import *

from apps.registration.models import User
#from apps.biblioteca.models import Modelo1
from docx.shared import RGBColor

modelo1_entradas = 0


def tablaNombramiento(document,request):

    # Variables contadoras de cada Categoria
    investigadorPosdoctoral = User.objects.filter(
        categoria='Investigador posdoctoral',coordinacion_id=request.user.coordinacion_id).count()
    catedraConacyt = User.objects.filter(categoria='Cátedra CONACyT',coordinacion_id=request.user.coordinacion_id).count()
    investigadorAsociadoC = User.objects.filter(
        categoria='Investigador Asociado C',coordinacion_id=request.user.coordinacion_id).count()
    investigadorTitularA = User.objects.filter(
        categoria='Investigador Titular A',coordinacion_id=request.user.coordinacion_id).count()
    investigadorTitularB = User.objects.filter(
        categoria='Investigador Titular B',coordinacion_id=request.user.coordinacion_id).count()
    investigadorTitularC = User.objects.filter(
        categoria='Investigador Titular C',coordinacion_id=request.user.coordinacion_id).count()
    investigadorTitularD = User.objects.filter(
        categoria='Investigador Titular D',coordinacion_id=request.user.coordinacion_id).count()
    TotalNombramientos = investigadorPosdoctoral + catedraConacyt + investigadorAsociadoC + \
        investigadorTitularA + investigadorTitularB + \
        investigadorTitularC + investigadorTitularD

    table = document.add_table(
        rows=1, cols=2, style='Medium Shading 1 Accent 1')

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nombramiento'
    hdr_cells[1].text = 'Número'

    row_cells = table.add_row().cells
    row_cells[0].text = "Postdocs"
    row_cells[1].text = str(investigadorPosdoctoral)

    row_cells = table.add_row().cells
    row_cells[0].text = "Cátedra CONACyT"
    row_cells[1].text = str(catedraConacyt)

    row_cells = table.add_row().cells
    row_cells[0].text = "Asociado C"
    row_cells[1].text = str(investigadorAsociadoC)

    row_cells = table.add_row().cells
    row_cells[0].text = "Investigador Titular A"
    row_cells[1].text = str(investigadorTitularA)

    row_cells = table.add_row().cells
    row_cells[0].text = "Investigador Titular B"
    row_cells[1].text = str(investigadorTitularB)

    row_cells = table.add_row().cells
    row_cells[0].text = "Investigador Titular C"
    row_cells[1].text = str(investigadorTitularC)

    row_cells = table.add_row().cells
    row_cells[0].text = "Investigador Titular D"
    row_cells[1].text = str(investigadorTitularD)

    row_cells = table.add_row().cells
    row_cells[0].text = "Total"
    row_cells[1].text = str(TotalNombramientos)
    document.add_paragraph()


def tablaNivelSNI(document,request):
    # Variables contadoras del Nivel SNI:
    sniCandidato = User.objects.filter(nivelSni='Candidato',coordinacion_id=request.user.coordinacion_id).count()
    sniNivel1 = User.objects.filter(nivelSni='Nivel 1',coordinacion_id=request.user.coordinacion_id).count()
    sniNivel2 = User.objects.filter(nivelSni='Nivel 2',coordinacion_id=request.user.coordinacion_id).count()
    sniNivel3 = User.objects.filter(nivelSni='Nivel 3',coordinacion_id=request.user.coordinacion_id).count()
    sniEmerito = User.objects.filter(nivelSni='Emérito',coordinacion_id=request.user.coordinacion_id).count()
    TotalSNI = sniCandidato + sniNivel1 + sniNivel2 + sniNivel3 + sniEmerito

    table = document.add_table(
        rows=1, cols=2, style='Medium Shading 1 Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # ALINEACION CENTRAL
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Distinción SNI'
    hdr_cells[1].text = 'Número'

    row_cells = table.add_row().cells
    row_cells[0].text = "Candidato"
    row_cells[1].text = str(sniCandidato)

    row_cells = table.add_row().cells
    row_cells[0].text = "Nivel 1"
    row_cells[1].text = str(sniNivel1)

    row_cells = table.add_row().cells
    row_cells[0].text = "Nivel 2"
    row_cells[1].text = str(sniNivel2)

    row_cells = table.add_row().cells
    row_cells[0].text = "Nivel 3"
    row_cells[1].text = str(sniNivel3)

    row_cells = table.add_row().cells
    row_cells[0].text = "Emérito"
    row_cells[1].text = str(sniEmerito)

    row_cells = table.add_row().cells
    row_cells[0].text = "Total"
    row_cells[1].text = str(TotalSNI)

    document.add_paragraph()


def tablaInvestigadores(document, usuario, reportes):
    table = document.add_table(
        rows=1, cols=5, style='Medium Shading 1 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Apellido'
    hdr_cells[1].text = 'Nombre'
    hdr_cells[2].text = 'Categoria'
    hdr_cells[3].text = 'SNI'
    hdr_cells[4].text = 'Reporte enviado'

    for r in reportes:
        row_cells = table.add_row().cells
        row_cells[0].text = r.usuario.apellido
        row_cells[1].text = r.usuario.nombre
        row_cells[2].text = r.usuario.categoria
        row_cells[3].text = r.usuario.nivelSni
        row_cells[4].text = 'ok'

    document.add_page_break()


def tablaResumenNumerales(document, yearPeriodo,request):
    paragraph = document.add_paragraph()
    paragraph.add_run("I.Resumen: Investigación").bold = True
    table = document.add_table(
        rows=1, cols=3, style='Medium Shading 1 Accent 1')

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Numeral'
    hdr_cells[1].text = 'Concepto'
    hdr_cells[2].text = 'Total en el periodo'

    row_cells = table.add_row().cells
    row_cells[0].text = "1"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas indizadas en el primer cuartil"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 1,'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "2"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas indizadas en segundo o tercer cuartil."
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 2,'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "3"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas indizadas en cuarto cuartil"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 3,'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "4"
    row_cells[1].text = "Artículos científicos arbitrados en revistas del Índice CONACYT"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 4,'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "5"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas emergentes"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 5,'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "6"
    row_cells[1].text = "Artículos científicos arbitrados en revistas periódicas no indizadas"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 6, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "7"
    row_cells[1].text = "Artículos aceptados con arbitraje internacional en revistas periódicas indizadas"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 7, 'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "8"
    row_cells[1].text = "Artículos aceptados con arbitraje en revistas periódicas no indizadas"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 8, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "9"
    row_cells[1].text = "Artículos enviados con arbitraje internacional en revistas periódicas indizadas"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 9, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "10"
    row_cells[1].text = "Artículos enviados con arbitraje en revistas periódicas no indizadas"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 10, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "11"
    row_cells[1].text = "Artículos científicos arbitrados en extenso en memorias de congresos internacionales"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 11, 'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "12"
    row_cells[1].text = "Artículos científicos arbitrados en extenso en memorias de congresos nacionales"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 12, 'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "13"
    row_cells[1].text = "Artículos científicos no arbitrados en extenso en memorias de congresos internacionales"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 13, 'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "14"
    row_cells[1].text = "Artículos científicos no arbitrados en extenso en memorias de congresos nacionales"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 14, 'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "14a"
    row_cells[1].text = "Reportes científicos no arbitrados en publicaciones periódicas."
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 14.1, 'bibcode',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "15"
    row_cells[1].text = "Autor o coautor de libros (no memorias de congreso)"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 15, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "16"
    row_cells[1].text = "Autor de capítulo de libro (no del mismo libro y no memoria de congreso)"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 16, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "17"
    row_cells[1].text = "Edición de libros / memorias"
    row_cells[2].text = Modelo1.objects.entradasContador(yearPeriodo, 17, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "18"
    row_cells[1].text = "Proyectos CONACyT"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo, 18, 'nombre_del_proyecto',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "19"
    row_cells[1].text = "Proyectos institucionales"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo, 19,'nombre_del_proyecto',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "20"
    row_cells[1].text = "Proyectos externos"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo, 20,'nombre_del_proyecto',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "21"
    row_cells[1].text = "Proyectos interinstitucionales"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo, 21, 'nombre_del_proyecto',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "22"
    row_cells[1].text = "Proyectos comercializados"
    row_cells[2].text = Modelo2.objects.entradasContador(yearPeriodo, 22, 'nombre_del_proyecto',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "23"
    row_cells[1].text = "Participación en el comité científico de conferencias internacionales (Scientific Organizing Committee; Steering Committee; similares)"
    row_cells[2].text = Modelo3.objects.entradasContador(yearPeriodo, 23, 'conferencia_proyecto',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "24"
    row_cells[1].text = "Conferencias científicas internacionales."
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo, 24, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "25"
    row_cells[1].text = "Conferencias científicas nacionales"
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo, 25, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "26"
    row_cells[1].text = "Pláticas invitadas en conferencias internacionales"
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo, 26, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "27"
    row_cells[1].text = "Pláticas invitadas en conferencias nacionales"
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo, 27, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "28"
    row_cells[1].text = "Resúmenes en congreso internacionales"
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo, 28, 'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "29"
    row_cells[1].text = "Resúmenes en congreso nacionales"
    row_cells[2].text = Modelo4.objects.entradasContador(yearPeriodo, 29, 'titulo',request)

    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.add_run("II.	Resumen: Formación de recursos humanos").bold = True
    table = document.add_table(
        rows=1, cols=3, style='Medium Shading 1 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Numeral'
    hdr_cells[1].text = 'Concepto'
    hdr_cells[2].text = 'Total en el periodo'

    row_cells = table.add_row().cells
    row_cells[0].text = "31"
    row_cells[1].text = "Alumnos graduados de doctorado en tiempos PNPC"
    row_cells[2].text = Modelo5.objects.entradasContador(yearPeriodo, 31,'nombre_completo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "32"
    row_cells[1].text = "Alumnos graduados de doctorado fuera de tiempo PNPC"
    row_cells[2].text = Modelo5.objects.entradasContador(yearPeriodo, 32,'nombre_completo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "33"
    row_cells[1].text = "Alumnos graduados de maestría en tiempos PNPC"
    row_cells[2].text = Modelo5.objects.entradasContador(yearPeriodo, 33,'nombre_completo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "34"
    row_cells[1].text = "Alumnos graduados de maestría fuera de tiempo PNPC"
    row_cells[2].text = Modelo5.objects.entradasContador(yearPeriodo, 34,'nombre_completo',request)
    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.add_run(
        "III.	Resumen: Desarrollo tecnológico e innovación").bold = True
    table = document.add_table(
        rows=1, cols=3, style='Medium Shading 1 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Numeral'
    hdr_cells[1].text = 'Concepto'
    hdr_cells[2].text = 'Total en el periodo'

    row_cells = table.add_row().cells  # Modelo 7  40-41-42-43-44
    row_cells[0].text = "40"
    row_cells[1].text = "Derechos de autor y aseguramiento de propiedad intelectual"
    row_cells[2].text = Modelo8.objects.entradasContador(yearPeriodo, 40,'autores',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "41"
    row_cells[1].text = "Patentes solicitadas"
    row_cells[2].text = Modelo8.objects.entradasContador(yearPeriodo, 41,'autores',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "42"
    row_cells[
        1].text = "Patentes en proceso de evaluación que ya aprobaron el examen de forma (IMPI)"
    row_cells[2].text = Modelo8.objects.entradasContador(yearPeriodo, 42,'autores',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "44"
    row_cells[1].text = "Patentes licenciadas"
    row_cells[2].text = Modelo8.objects.entradasContador(yearPeriodo, 44,'autores',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "45"
    row_cells[1].text = "Dirección de proyectos de investigación tecnológica"
    row_cells[2].text = Modelo9.objects.entradasContador(yearPeriodo, 45,'nombre',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "46"
    row_cells[1].text = "Reportes técnicos registrados"
    row_cells[2].text = Modelo10.objects.entradasContador(yearPeriodo, 46, 'titulo',request)

    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.add_run("IV. Resumen: Apoyo Institucional").bold = True
    table = document.add_table(
        rows=1, cols=3, style='Medium Shading 1 Accent 1')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Numeral'
    hdr_cells[1].text = 'Concepto'
    hdr_cells[2].text = 'Total en el periodo'

    row_cells = table.add_row().cells
    row_cells[0].text = "48"
    row_cells[1].text = "Artículos de divulgación científica en medios masivos"
    row_cells[2].text = Modelo10.objects.entradasContador(yearPeriodo, 48,'titulo',request)

    row_cells = table.add_row().cells
    row_cells[0].text = "49"
    row_cells[1].text = "Conferencias de divulgación en eventos masivos"
    row_cells[2].text = Modelo11.objects.entradasContador(yearPeriodo, 49,request)

    row_cells = table.add_row().cells
    row_cells[0].text = "50"
    row_cells[1].text = "Conferencias de difusión o promoción externas"
    row_cells[2].text = Modelo11.objects.entradasContador(yearPeriodo, 50,request)

    row_cells = table.add_row().cells
    row_cells[0].text = "51"
    row_cells[1].text = "Conferencias de difusión o promoción internas"
    row_cells[2].text = Modelo11.objects.entradasContador(yearPeriodo, 51,request)

    row_cells = table.add_row().cells
    row_cells[0].text = "52"
    row_cells[1].text = "Organización de eventos académicos vinculados al quehacer institucional"
    row_cells[2].text = Modelo11.objects.entradasContador(yearPeriodo, 52,request)


def filtroAutores(cadenaAutores, nombreC, paragraph):
    autores = cadenaAutores.split(";")

    for a in autores:
        if a.strip().split(",")[0] == nombreC.split(",")[0]:
            paragraph.add_run(a).font.bold = True
        else:
            paragraph.add_run(a).font.bold = False


def nombreFile(nombreFile):
    nombre = nombreFile.replace("anexo/", "")
    return nombre


def citas(document, reportes,request):
    citasObj = Citas.objects.filter(usuario__coordinacion=request.user.coordinacion)
    # Tabla Distincion SNI
    table = document.add_table(
        rows=1, cols=4, style='Medium Shading 1 Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # ALINEACION CENTRAL
    table.allow_autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Investigador'
    hdr_cells[1].text = 'Citas totales'
    hdr_cells[2].text = 'Citas en el periodo'
    hdr_cells[3].text = 'Índice-h'
    for r in reportes:
        for item in citasObj:
            if r.usuario_id == item.usuario_id:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.usuario.nombreCorto)
                row_cells[1].text = str(item.citas)
                row_cells[2].text = str(item.citas_en_periodo)
                row_cells[3].text = str(item.indiceH)
    document.add_paragraph()


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def Reporte(request, periodo_id):
    print(request.user)
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    nombreAux = True
    # Objetos
    # Para obtener valores del 2019
    periodo = Periodo.objects.get(id=periodo_id)
    yearPeriodo = periodo.fecha_inicio.year
    fecha_inicioPeriodo = periodo.fecha_inicio
    biblioteca = Modelo1.objects.filter(fecha_ano=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id)
    reportes = ReporteEnviado.objects.filter(periodo=periodo,usuario__coordinacion_id=request.user.coordinacion_id)
    usuario = User.objects.filter(is_staff=0,coordinacion_id=request.user.coordinacion_id).order_by('apellido')  # exclude(email = 'astrofi@inaoep.mx')
    modelo1 = Modelo1.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('titulo').distinct('titulo')
    modelo2 = Modelo2.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('nombre_del_proyecto').distinct('nombre_del_proyecto')
    modelo3 = Modelo3.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('conferencia_proyecto').distinct('conferencia_proyecto')
    modelo4 = Modelo4.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('titulo').distinct('titulo')
    modelo5 = Modelo5.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('nombre_completo').distinct('nombre_completo')
    modelo6 = Modelo6.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('nombre_del_curso').distinct('nombre_del_curso')
    modelo7 = Modelo7.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('nombre').distinct('nombre')
    modelo8 = Modelo8.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('autores').distinct('autores')
    modelo9 = Modelo9.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('nombre').distinct('nombre')
    modelo10 = Modelo10.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id).order_by('titulo').distinct('titulo')
    modelo11 = Modelo11.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id)
    modelo12 = Modelo12.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id)
    modelo13 = Modelo13.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id)
    modelo14 = Modelo14.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id)
    modelo15 = Modelo15.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id)
    modelo16 = Modelo16.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id)
    modelo17 = Modelo17.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id)
    modelo18 = Modelo18.objects.filter(periodo__fecha_inicio__year=yearPeriodo,usuario__coordinacion_id=request.user.coordinacion_id)
    numeral = Numeral.objects.all()
    # Fin objetos

    if fecha_inicioPeriodo.month == 1:
        MesInicio = "ENERO"
        MesFin = "JUNIO"
    elif fecha_inicioPeriodo.month == 6:
        MesInicio = "JUNIO"
        MesFin = "DICIEMBRE"

    # Titulos del documento
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = paragraph.add_run()
    p.bold = True
    p.text = "INAOE"
    p.font.size = Pt(13)

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = paragraph.add_run()
    p.bold = True
    p.text = "Coordinación de "+request.user.coordinacion.nombre
    p.font.size = Pt(13)

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.underline = True
    p = paragraph.add_run()
    p.bold = True
    p.text = "Reporte de Productividad Unificado"
    p.font.size = Pt(13)

    paragraph = document.add_paragraph()
    p = paragraph.add_run()
    p.text = "PERIODO DEL REPORTE: "+MesInicio + \
        "-"+MesFin+" DE " + str(yearPeriodo)
    p.bold = True
    p.underline = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Fin titulos documento

    document.add_paragraph()
    paragraph = document.add_paragraph()
    p = paragraph.add_run()
    p.bold = True
    p.text = "Estadísticas de reportes: Coordinación de "+request.user.coordinacion.nombre

    paragraph = document.add_paragraph(
        "Número total de investigadores en la plataforma de reportes: ")
    p = paragraph.add_run()
    p.bold = True
    p.text = str(usuario.count())

    # Tabla 1 de nombramiento
    tablaNombramiento(document,request)

    # Tabla 2 de Distincion SNI
    tablaNivelSNI(document,request)

    # Tabla 3 Datos reportes
    tablaInvestigadores(document, usuario, reportes)

    # Tabla 4 resumen Numerales
    tablaResumenNumerales(document, yearPeriodo,request)

    document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph.add_run("I. INVESTIGACIÓN CIENTÍFICA").bold = True

    for n in numeral:
        # Secciones
        if n.orden == 31:
            paragraph = document.add_paragraph()
            paragraph.add_run("II. FORMACIÓN DE RECURSOS HUMANOS").bold = True

        if n.orden == 40:
            document.add_paragraph()
            paragraph = document.add_paragraph()
            paragraph.add_run(
                "III. DESARROLLO TECNOLÓGICO E INNOVACIÓN").bold = True

        if n.orden == 47:
            document.add_paragraph()
            paragraph = document.add_paragraph()
            paragraph.add_run("IV. APOYO INSTITUCIONAL").bold = True

        if n.orden == 61:
            document.add_paragraph()
            paragraph = document.add_paragraph()
            paragraph.add_run("V. INFORMACIÓN ADICIONAL").bold = True

        document.add_paragraph(n.nombre)  # Nombre - Numerales
        if n.orden == 30:
            citas(document, reportes,request)

        for r in reportes:
            nombreAux = True

            for item in modelo1:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    filtroAutores(
                        item.autores, r.usuario.nombreCorto, paragraph)
                    paragraph.add_run(", ")

                    paragraph.add_run(item.titulo).font.italic = True
                    paragraph.add_run(", ")

                    if item.revista_publicacion:
                        paragraph.add_run(item.revista_publicacion)
                        paragraph.add_run(" ")
                    
                    if item.volumen:
                        paragraph.add_run(item.volumen)
                        paragraph.add_run(" ")

                    if item.paginas:
                        paragraph.add_run(item.paginas)
                        paragraph.add_run(" ")

                    paragraph.add_run(item.fecha)

                    p = document.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.left_indent = Inches(0.25)

                    if item.estudiantes:
                        p.add_run("Estudiante(s): ")
                        p.add_run(item.estudiantes).font.color.rgb = RGBColor(
                            255, 0, 0)
                        p.add_run(" \r")

                    if item.doi:
                        p.add_run("DOI: ")
                        p.add_run(item.doi).font.color.rgb = RGBColor(
                            0x42, 0x24, 0xE9)
                        p.add_run(" \r")
                    
                    if item.bibcode:
                        p.add_run("Bibcode: "+item.bibcode)
                        p.add_run(" \r")

                    if item.url:
                        p.add_run("URL: ")
                        add_hyperlink(p, item.url, item.url)
                        p.add_run(" \r")

                    if item.anexo:
                        p.add_run("Anexo: " + nombreFile(item.anexo.name))

            for item in modelo2:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.nombre_del_proyecto:
                        paragraph.add_run(
                            "Nombre del proyecto: " + item.nombre_del_proyecto + "\n")

                    if item.participantes:
                        paragraph.add_run(
                            "Participantes: " + item.participantes + "\n")

                    if item.rol:
                        paragraph.add_run("Rol: " + item.rol + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripcion: " +
                                          item.descripcion + "\n")

                    if item.estudiantes:
                        paragraph.add_run("Estudiantes: ")
                        paragraph.add_run(
                            item.estudiantes + "\n").font.color.rgb = RGBColor(255, 0, 0)

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo3:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.conferencia_proyecto:
                        paragraph.add_run(
                            "Conferencia o proyecto: " + item.conferencia_proyecto + "\n")

                    if item.rol:
                        paragraph.add_run("Rol: " + item.rol)

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph, item.url, item.url)
                        paragraph.add_run(" \r")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo4:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.titulo:
                        paragraph.add_run(
                            "Título de la presentación: " + item.titulo + "\n")

                    if item.autores:
                        paragraph.add_run("Autor(es): " + item.autores + "\n")

                    if item.nombre_de_conferencia:
                        paragraph.add_run(
                            "Nombre de la conferencia: " + item.nombre_de_conferencia + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.tipo:
                        paragraph.add_run(
                            "Tipo de presentación: " + item.tipo + "\n")

                    if item.estudiantes:
                        paragraph.add_run("Estudiantes: ")
                        paragraph.add_run(
                            item.estudiantes + "\n").font.color.rgb = RGBColor(255, 0, 0)

                    if item.doi:
                        paragraph.add_run("DOI/ISBN: " + item.doi + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph, item.url, item.url)
                        paragraph.add_run(" \r")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo5:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.nombre_completo:
                        paragraph.add_run(
                            "Nombre completo: " + item.nombre_completo + "\n")

                    if item.titulo_de_tesis:
                        paragraph.add_run(
                            "Título de tesis: " + item.titulo_de_tesis + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph, item.url, item.url)
                        paragraph.add_run(" \r")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo6:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.nombre_del_curso:
                        paragraph.add_run(
                            "Nombre del curso: " + item.nombre_del_curso + "\n")

                    if item.periodo_numeral:
                        paragraph.add_run(
                            "Periodo: " + item.periodo_numeral + "\n")

                    if item.notas:
                        paragraph.add_run("Notas: " + item.notas + "\n")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo7:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.nombre:
                        paragraph.add_run("Nombre: " + item.nombre + "\n")

                    if item.titulo_de_tesis:
                        paragraph.add_run(
                            "Título de tesis: " + item.titulo_de_tesis + "\n")

                    if item.grado:
                        paragraph.add_run("Grado: " + item.grado + "\n")

                    if item.institucion:
                        paragraph.add_run("Institución: " +
                                          item.institucion + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.notas:
                        paragraph.add_run("Notas: " + item.notas + "\n")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo8:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.autores:
                        paragraph.add_run("Autor(es): " + item.autores + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " +
                                          item.descripcion + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph, item.url, item.url)
                        paragraph.add_run(" \r")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo9:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.nombre:
                        paragraph.add_run("Nombre: " + item.nombre + "\n")

                    if item.participantes:
                        paragraph.add_run(
                            "Participantes: " + item.participantes + "\n")

                    if item.financiamiento:
                        paragraph.add_run(
                            "Financiamiento: " + item.financiamiento + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " +
                                          item.descripcion + "\n")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo10:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.titulo:
                        paragraph.add_run("Titulo: " + item.titulo + "\n")

                    if item.autores:
                        paragraph.add_run("Autor(es): " + item.autores + "\n")

                    if item.numero_reportes:
                        paragraph.add_run(
                            "No. Reporte/ID: " + item.numero_reportes + "\n")

                    if item.revista_publicacion:
                        paragraph.add_run(
                            "Revista o publicación: " + item.revista_publicacion + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.doi:
                        paragraph.add_run("DOI/ISBN: " + item.doi + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph, item.url, item.url)
                        paragraph.add_run(" \r")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo11:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph, item.url, item.url + "\n")
                        paragraph.add_run(" \r")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " +
                                          item.descripcion + "\n")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo12:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.periodo_numeral:
                        paragraph.add_run("Periodo: " + item.periodo_numeral)

                    if item.descripcion:
                        paragraph.add_run("Descripción: " +
                                          item.descripcion + "\n")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo13:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.nombre_del_estudiante:
                        paragraph.add_run(
                            "Nombre del estudiante: " + item.nombre_del_estudiante + "\n")

                    if item.fecha:
                        paragraph.add_run("Fecha: " + item.fecha + "\n")

                    if item.coordinacion:
                        paragraph.add_run("Coordinación: " + item.coordinacion)

                    if item.grado:
                        paragraph.add_run("Grado: " + item.grado)

                    if item.descripcion:
                        paragraph.add_run("Descripción: " +
                                          item.descripcion + "\n")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo14:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.fecha_periodo:
                        paragraph.add_run("Fecha o periodo: " + item.fecha_periodo + '\n')

                    if item.descripcion:
                        paragraph.add_run("Descripción: " + '\n')

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo15:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.laboratorio_taller:
                        paragraph.add_run(
                            "Laboratorio o taller:" + item.laboratorio_taller + "\n")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo16:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style='List Bullet')

                    if item.agencias_financieras:
                        paragraph.add_run(
                            "Agencia(s) financiadora(s):" + item.agencias_financieras + "\n")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " +
                                          item.descripcion + "\n")

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo17:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style="List Bullet")

                    if item.telescopio_instrumento_infra:
                        paragraph.add_run(
                            "Telescopio, instrumento, infraestructura: " + item.telescopio_instrumento_infra + "\n")

                    if item.url:
                        paragraph.add_run("URL: ")
                        add_hyperlink(paragraph, item.url, item.url + "\n")
                        paragraph.add_run(" \r")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion)

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

            for item in modelo18:
                if n.id == item.numeral_id and r.usuario_id == item.usuario_id:
                    if(nombreAux):
                        paragraph = document.add_paragraph()
                        paragraph.add_run(
                            r.usuario.nombreCorto).font.bold = True
                    nombreAux = False
                    paragraph = document.add_paragraph(style="List Bullet")

                    if item.descripcion:
                        paragraph.add_run("Descripción: " + item.descripcion)

                    if item.anexo:
                        paragraph.add_run(
                            "Anexo: " + nombreFile(item.anexo.name))

    return document
